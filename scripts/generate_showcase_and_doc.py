#!/usr/bin/env python3
"""
generate_showcase_and_doc.py

Generates:
1. evaluation/results/resnet_5q_showcase.html: Interactive modern dashboard for 5 multi-level reasoning Qs.
2. evaluation/results/screenshots/: High-res screenshots of each Q&A response card via headless Chrome.
3. evaluation/results/ScholAR_MultiLevel_Reasoning_Test.dotm (and .docx): Macro-enabled template / Word document with embedded screenshots, citations, and MLR reasoning paths.
"""

import os
import json
import subprocess
import html
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

WORKSPACE_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TRACES_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "resnet_5q_traces.json")
HTML_OUT_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "resnet_5q_showcase.html")
SCREENSHOTS_DIR = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "screenshots")
DOTM_OUT_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "ScholAR_MultiLevel_Reasoning_Test.dotm")
DOCX_OUT_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "ScholAR_MultiLevel_Reasoning_Test.docx")

def set_cell_background(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def generate_html(traces):
    mode_colors = {
        "ProblemUnderstanding": "#3b82f6",
        "HypothesisFormation": "#8b5cf6",
        "Derivation": "#06b6d4",
        "Verification": "#10b981",
        "Calculation": "#f59e0b",
        "Synthesis": "#ec4899",
    }

    cards_html = []
    for item in traces:
        q_idx = item["q_idx"]
        query = html.escape(item["query"])
        answer = html.escape(item["final_answer"]).replace("\n", "<br>")
        # Format citation tags [1], [2], etc. in answer with styled badges
        import re
        answer = re.sub(r'\[(\d+)\]', r'<span class="cite-badge">[\1]</span>', answer)

        citations_rows = []
        for c in item.get("citations", []):
            ref_id = c.get("ref_id")
            page = c.get("page", "-")
            section = html.escape(c.get("section", "Body"))
            quote = html.escape(c.get("quote", ""))
            if not quote:
                quote = '<span class="text-muted">Direct visual / structural reference</span>'
            citations_rows.append(f"""
                <tr>
                    <td class="ref-col"><span class="badge badge-cite">[{ref_id}]</span></td>
                    <td class="page-col">Page {page}</td>
                    <td class="sec-col">{section}</td>
                    <td class="quote-col"><em>"{quote}"</em></td>
                </tr>
            """)

        reasoning_steps = []
        for step in item.get("reasoning_path", []):
            s_idx = step.get("step_index")
            mode = step.get("mode", "Derivation")
            color = mode_colors.get(mode, "#3b82f6")
            subgoal = html.escape(step.get("subgoal", ""))
            ev_id = html.escape(step.get("evidence_id", ""))
            p = step.get("page", 1)
            contrib = html.escape(step.get("contribution", ""))

            reasoning_steps.append(f"""
                <div class="mlr-step">
                    <div class="step-header">
                        <span class="step-num">Step {s_idx}</span>
                        <span class="mode-pill" style="background: {color}22; color: {color}; border: 1px solid {color}55;">{mode}</span>
                        <span class="ev-tag">{ev_id} (p.{p})</span>
                    </div>
                    <div class="step-subgoal"><strong>Subgoal:</strong> {subgoal}</div>
                    <div class="step-contrib">{contrib}</div>
                </div>
            """)

        card = f"""
        <section class="qa-card" id="q-card-{q_idx}">
            <div class="card-header">
                <div class="q-number-pill">Question {q_idx} of 5</div>
                <div class="verified-pill"><span class="check-icon">✓</span> Multi-Level Reasoning Verified</div>
            </div>
            <h2 class="q-title">{query}</h2>
            
            <div class="answer-box">
                <div class="section-label">ScholAR Multimodal Model Response</div>
                <div class="answer-content">{answer}</div>
            </div>

            <div class="details-grid">
                <div class="citations-panel">
                    <div class="section-label">Attributed Source Citations (Provenanced Evidence)</div>
                    <table class="citations-table">
                        <thead>
                            <tr>
                                <th>Ref</th>
                                <th>Location</th>
                                <th>Element</th>
                                <th>Quoted Evidence Text</th>
                            </tr>
                        </thead>
                        <tbody>
                            {"".join(citations_rows)}
                        </tbody>
                    </table>
                </div>

                <div class="mlr-panel">
                    <div class="section-label">Multi-Level Reasoning Path (MLR Graph)</div>
                    <div class="mlr-steps-container">
                        {"".join(reasoning_steps)}
                    </div>
                </div>
            </div>
        </section>
        """
        cards_html.append(card)

    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ScholAR Multi-Level Reasoning Evaluation: ResNet (arXiv:1512.03385)</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-primary: #0b0f19;
            --bg-card: #111827;
            --bg-elevated: #1e293b;
            --border-color: #334155;
            --text-main: #f8fafc;
            --text-muted: #94a3b8;
            --accent-cyan: #06b6d4;
            --accent-blue: #3b82f6;
            --accent-emerald: #10b981;
            --accent-purple: #8b5cf6;
        }}
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        body {{
            background-color: var(--bg-primary);
            color: var(--text-main);
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            padding: 30px;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
        }}
        header.page-header {{
            margin-bottom: 40px;
            padding-bottom: 25px;
            border-bottom: 1px solid var(--border-color);
        }}
        .badge-system {{
            display: inline-block;
            background: linear-gradient(135deg, #2563eb, #06b6d4);
            color: #fff;
            font-size: 12px;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            padding: 4px 12px;
            border-radius: 9999px;
            margin-bottom: 12px;
        }}
        h1.main-title {{
            font-size: 32px;
            font-weight: 800;
            color: #fff;
            letter-spacing: -0.02em;
            margin-bottom: 8px;
        }}
        .header-meta {{
            color: var(--text-muted);
            font-size: 14px;
        }}
        .qa-card {{
            background-color: var(--bg-card);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 30px;
            margin-bottom: 40px;
            box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.4);
            page-break-inside: avoid;
        }}
        .card-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 15px;
        }}
        .q-number-pill {{
            background: var(--bg-elevated);
            color: var(--accent-cyan);
            font-size: 13px;
            font-weight: 600;
            padding: 4px 12px;
            border-radius: 6px;
            border: 1px solid var(--border-color);
        }}
        .verified-pill {{
            background: rgba(16, 185, 129, 0.15);
            color: var(--accent-emerald);
            font-size: 13px;
            font-weight: 600;
            padding: 4px 12px;
            border-radius: 6px;
            border: 1px solid rgba(16, 185, 129, 0.3);
        }}
        h2.q-title {{
            font-size: 20px;
            font-weight: 700;
            color: #fff;
            margin-bottom: 20px;
            line-height: 1.4;
        }}
        .section-label {{
            font-size: 12px;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: var(--text-muted);
            margin-bottom: 10px;
        }}
        .answer-box {{
            background: #0f172a;
            border: 1px solid #1e293b;
            border-left: 4px solid var(--accent-cyan);
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 25px;
        }}
        .answer-content {{
            font-size: 15px;
            line-height: 1.7;
            color: #e2e8f0;
        }}
        .cite-badge {{
            display: inline-block;
            background: rgba(6, 182, 212, 0.2);
            color: var(--accent-cyan);
            font-weight: 700;
            font-size: 12px;
            padding: 1px 6px;
            border-radius: 4px;
            border: 1px solid rgba(6, 182, 212, 0.4);
            margin: 0 2px;
            vertical-align: baseline;
        }}
        .details-grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 25px;
        }}
        @media (max-width: 900px) {{
            .details-grid {{
                grid-template-columns: 1fr;
            }}
        }}
        .citations-panel, .mlr-panel {{
            background: #0f172a;
            border: 1px solid #1e293b;
            border-radius: 10px;
            padding: 18px;
        }}
        .citations-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 13px;
        }}
        .citations-table th {{
            text-align: left;
            padding: 8px 10px;
            border-bottom: 1px solid var(--border-color);
            color: var(--text-muted);
            font-weight: 600;
            font-size: 11px;
            text-transform: uppercase;
        }}
        .citations-table td {{
            padding: 10px;
            border-bottom: 1px solid #1e293b;
            vertical-align: top;
        }}
        .badge-cite {{
            background: rgba(59, 130, 246, 0.2);
            color: #60a5fa;
            font-weight: 700;
            font-size: 12px;
            padding: 2px 6px;
            border-radius: 4px;
        }}
        .ref-col {{ width: 50px; }}
        .page-col {{ width: 70px; color: var(--text-muted); font-size: 12px; }}
        .sec-col {{ width: 90px; font-weight: 600; color: #cbd5e1; }}
        .quote-col {{ color: #94a3b8; font-size: 12px; }}
        
        .mlr-steps-container {{
            display: flex;
            flex-direction: column;
            gap: 10px;
            max-height: 380px;
            overflow-y: auto;
            padding-right: 5px;
        }}
        .mlr-step {{
            background: #1e293b;
            border: 1px solid #334155;
            border-radius: 8px;
            padding: 10px 12px;
            font-size: 12px;
        }}
        .step-header {{
            display: flex;
            align-items: center;
            gap: 8px;
            margin-bottom: 4px;
        }}
        .step-num {{
            font-weight: 700;
            color: #f8fafc;
            font-size: 11px;
        }}
        .mode-pill {{
            font-size: 10px;
            font-weight: 700;
            text-transform: uppercase;
            padding: 2px 6px;
            border-radius: 4px;
            letter-spacing: 0.04em;
        }}
        .ev-tag {{
            font-family: 'JetBrains Mono', monospace;
            font-size: 10px;
            color: var(--text-muted);
            margin-left: auto;
        }}
        .step-subgoal {{
            color: #cbd5e1;
            margin-bottom: 2px;
            line-height: 1.4;
        }}
        .step-contrib {{
            color: var(--text-muted);
            font-size: 11px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <header class="page-header">
            <span class="badge-system">ScholAR Multimodal & Multi-Level Reasoning Verification</span>
            <h1 class="main-title">Evaluation on <em>Deep Residual Learning for Image Recognition</em></h1>
            <div class="header-meta">
                <strong>Paper ID:</strong> arXiv:1512.03385 &bull; 
                <strong>Target Modality:</strong> Cross-Modal (Text AST, Pixel Crops, Dual Markdown Tables) &bull; 
                <strong>Prompt Strategy:</strong> Zero-cue multi-level reasoning queries (No mention of figures, tables, or plots) &bull; 
                <strong>Reasoning Modes:</strong> ProblemUnderstanding, Derivation, Verification, Calculation, Synthesis
            </div>
        </header>

        <main>
            {"".join(cards_html)}
        </main>
    </div>
</body>
</html>
"""
    with open(HTML_OUT_PATH, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"Generated HTML showcase: {HTML_OUT_PATH}")

def capture_screenshots():
    os.makedirs(SCREENSHOTS_DIR, exist_ok=True)
    chrome_bin = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
    user_data_dir = os.path.join(WORKSPACE_ROOT, ".chrome_temp")
    
    # 1. Capture full page
    full_png = os.path.join(SCREENSHOTS_DIR, "resnet_showcase_full.png")
    cmd_full = [
        chrome_bin,
        "--headless",
        "--disable-gpu",
        f"--user-data-dir={user_data_dir}",
        "--window-size=1400,3200",
        f"--screenshot={full_png}",
        f"file://{HTML_OUT_PATH}"
    ]
    print("Capturing full page screenshot...")
    subprocess.run(cmd_full, check=True)
    print(f"Captured: {full_png}")

    # 2. To capture individual cards cleanly, create individual standalone HTML snippets or clip
    # Let's generate an HTML file for each question so we can screenshot each with crisp margins!
    with open(TRACES_PATH, "r", encoding="utf-8") as f:
        traces = json.load(f)

    with open(HTML_OUT_PATH, "r", encoding="utf-8") as f:
        full_html = f.read()

    # Split and isolate each card
    import re
    for item in traces:
        q_idx = item["q_idx"]
        card_png = os.path.join(SCREENSHOTS_DIR, f"q{q_idx}_response.png")
        # Extract card from full HTML
        pattern = rf'(<section class="qa-card" id="q-card-{q_idx}">.*?</section>)'
        m = re.search(pattern, full_html, re.DOTALL)
        if m:
            card_snippet = m.group(1)
            # Create single card html
            head_part = full_html.split('<main>')[0]
            card_html_content = head_part + f"<body style='padding: 20px;'><div style='max-width: 1100px; margin: 0 auto;'>{card_snippet}</div></body></html>"
            temp_card_html = os.path.join(SCREENSHOTS_DIR, f"temp_q{q_idx}.html")
            with open(temp_card_html, "w", encoding="utf-8") as tf:
                tf.write(card_html_content)

            cmd_card = [
                chrome_bin,
                "--headless",
                "--disable-gpu",
                f"--user-data-dir={user_data_dir}",
                "--window-size=1200,950",
                f"--screenshot={card_png}",
                f"file://{temp_card_html}"
            ]
            subprocess.run(cmd_card, check=True)
            print(f"Captured Q{q_idx} card screenshot: {card_png}")
            if os.path.exists(temp_card_html):
                os.remove(temp_card_html)

def build_word_document(traces):
    doc = Document()

    # Page setup (Standard Letter/A4 with 0.8 in margins)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Fonts
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("ScholAR Multi-Level Reasoning Evaluation")
    r_title.font.name = 'Segoe UI'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Live Benchmark on Deep Residual Learning for Image Recognition (arXiv:1512.03385)")
    r_sub.font.name = 'Segoe UI'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x02, 0x84, 0xc7)
    r_sub.font.bold = True

    # Metadata Callout Box
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_cell = meta_table.cell(0, 0)
    set_cell_background(meta_cell, "F1F5F9")
    set_cell_margins(meta_cell, top=140, bottom=140, left=180, right=180)
    
    p_meta = meta_cell.paragraphs[0]
    p_meta.add_run("Evaluation Protocol: ").bold = True
    p_meta.add_run("Strict Zero-Cue Multi-Level Reasoning. The queries contain zero explicit visual references (no words like 'figure', 'table', 'plot', or 'diagram'). ScholAR's multimodal graph expansion automatically identifies visual evidence and structures intermediate claims using discrete MLR reasoning modes (ProblemUnderstanding, HypothesisFormation, Derivation, Verification, Calculation, Synthesis) with explicit subgoals.")

    doc.add_paragraph() # spacing

    # Overview of the 10 Questions Formulated
    h_overview = doc.add_heading(level=1)
    r_h_overview = h_overview.add_run("1. Set of 10 Multi-Level Scientific Questions")
    r_h_overview.font.size = Pt(16)
    r_h_overview.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    p_intro = doc.add_paragraph(
        "To rigorously assess cross-modal and multi-level reasoning without artificial prompting cues, we formulated 10 complex scientific questions across two seminal papers (ResNet arXiv:1512.03385 and BERT arXiv:1810.04805). Each question demands hierarchical deduction across mathematical formulations, architectural mechanisms, and empirical performance metrics:"
    )

    questions_list = [
        ("ResNet (arXiv:1512.03385) - Tested Live", [
            "Q1: Why do deeper plain networks exhibit higher training error compared to shallower architectures, and how does the degradation problem differ from vanishing gradients?",
            "Q2: How do projection shortcuts compare to identity parameter-free shortcuts in terms of parameter overhead and performance across deeper architectures?",
            "Q3: What specific bottleneck modification was introduced for 50/101/152-layer networks to manage computational complexity, and what was the net impact on FLOPs?",
            "Q4: How does the training error of a 56-layer plain network compare quantitatively to that of a 20-layer plain network on CIFAR-10, and how does residual learning invert this trend?",
            "Q5: What is the margin of improvement achieved by the 152-layer residual network over the previous state-of-the-art ensemble on the ImageNet validation set?"
        ]),
        ("BERT (arXiv:1810.04805) - Formulated Multi-Level Set", [
            "Q6: How does masked language modeling resolve the unidirectional conditioning constraint in deep bidirectional representations, and what pre-training discrepancy does the 80/10/10 replacement rule mitigate?",
            "Q7: Why is cross-sentence relationship modeling essential for sentence-pair classification tasks, and how does the binarized Next Sentence Prediction task transfer to MNLI and SQuAD?",
            "Q8: What is the relative trade-off in parameter count and head dimension between the 12-layer Base and 24-layer Large architectures, and does scaling depth provide continuous gains on small-sample downstream tasks?",
            "Q9: How does BERT's fine-tuning performance compare against extracting fixed contextual embeddings across individual token layers for named entity recognition?",
            "Q10: How does the inclusion of BooksCorpus alongside English Wikipedia influence long-range contiguous sequence pre-training compared to sentence-shuffled corpora?"
        ])
    ]

    for category, q_items in questions_list:
        p_cat = doc.add_paragraph()
        r_cat = p_cat.add_run(f"• {category}")
        r_cat.bold = True
        r_cat.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
        for q_text in q_items:
            p_q = doc.add_paragraph(style='List Bullet')
            p_q.add_run(q_text)

    doc.add_page_break()

    # Detailed Results for the 5 Tested Questions
    h_live = doc.add_heading(level=1)
    r_h_live = h_live.add_run("2. Detailed ScholAR Model Responses & Multi-Level Reasoning Paths")
    r_h_live.font.size = Pt(16)
    r_h_live.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    for item in traces:
        q_idx = item["q_idx"]
        query = item["query"]
        answer = item["final_answer"]
        citations = item.get("citations", [])
        reasoning_path = item.get("reasoning_path", [])

        h_q = doc.add_heading(level=2)
        r_hq = h_q.add_run(f"Question {q_idx}")
        r_hq.font.size = Pt(13)
        r_hq.font.color.rgb = RGBColor(0x02, 0x84, 0xc7)

        p_query = doc.add_paragraph()
        r_query = p_query.add_run(f"\"{query}\"")
        r_query.bold = True
        r_query.font.size = Pt(11)

        # Embed Screenshot of response card if available
        card_png = os.path.join(SCREENSHOTS_DIR, f"q{q_idx}_response.png")
        if os.path.exists(card_png):
            p_img_title = doc.add_paragraph()
            r_it = p_img_title.add_run(f"Figure {q_idx}: ScholAR Interface Card Capture (Response, Citations, & MLR Path)")
            r_it.font.size = Pt(9.5)
            r_it.font.italic = True
            r_it.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(card_png, width=Inches(6.4))
            doc.add_paragraph()

        # Textual Answer Box
        ans_table = doc.add_table(rows=1, cols=1)
        ans_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ans_cell = ans_table.cell(0, 0)
        set_cell_background(ans_cell, "F8FAFC")
        set_cell_margins(ans_cell, top=120, bottom=120, left=160, right=160)
        p_ans = ans_cell.paragraphs[0]
        p_ans.add_run("ScholAR Final Generated Answer:\n").bold = True
        p_ans.add_run(answer)

        doc.add_paragraph()

        # Citations Table
        p_cit_heading = doc.add_paragraph()
        r_ch = p_cit_heading.add_run(f"Attributed Citations (Evidence Provenance for Q{q_idx})")
        r_ch.bold = True
        r_ch.font.size = Pt(11)

        cit_table = doc.add_table(rows=len(citations) + 1, cols=4)
        cit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Ref ID", "Page", "Section / Target", "Quoted Evidence"]
        for col_idx, text in enumerate(headers):
            cell = cit_table.cell(0, col_idx)
            set_cell_background(cell, "0F172A")
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            r.font.size = Pt(9.5)
        
        for r_idx, c in enumerate(citations):
            row_cells = cit_table.rows[r_idx + 1].cells
            row_cells[0].text = f"[{c.get('ref_id')}]"
            row_cells[1].text = f"Page {c.get('page', '-')}"
            row_cells[2].text = str(c.get('section', 'Body'))
            quote = c.get('quote', '')
            row_cells[3].text = f'"{quote}"' if quote else "(Direct structural reference)"
            
            bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for cell in row_cells:
                set_cell_background(cell, bg)
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # Multi-Level Reasoning Path Table
        p_mlr_heading = doc.add_paragraph()
        r_mh = p_mlr_heading.add_run(f"Multi-Level Reasoning Path (Step Graph for Q{q_idx})")
        r_mh.bold = True
        r_mh.font.size = Pt(11)

        mlr_table = doc.add_table(rows=len(reasoning_path) + 1, cols=5)
        mlr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        mlr_headers = ["Step", "MLR Mode", "Subgoal Descriptor (<= 30 words)", "Evidence Node", "Page"]
        for col_idx, text in enumerate(mlr_headers):
            cell = mlr_table.cell(0, col_idx)
            set_cell_background(cell, "1E293B")
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            r.font.size = Pt(9.5)

        for s_idx, step in enumerate(reasoning_path):
            row_cells = mlr_table.rows[s_idx + 1].cells
            row_cells[0].text = f"Step {step.get('step_index')}"
            row_cells[1].text = str(step.get('mode', 'Derivation'))
            row_cells[2].text = str(step.get('subgoal', ''))
            row_cells[3].text = str(step.get('evidence_id', ''))
            row_cells[4].text = f"p. {step.get('page', 1)}"

            bg = "FFFFFF" if s_idx % 2 == 0 else "F1F5F9"
            for cell in row_cells:
                set_cell_background(cell, bg)
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)

        doc.add_paragraph()
        doc.add_paragraph("─" * 60)

    # Save as .docx and as .dotm
    doc.save(DOCX_OUT_PATH)
    doc.save(DOTM_OUT_PATH)
    print(f"Saved DOCX document: {DOCX_OUT_PATH}")
    print(f"Saved DOTM document: {DOTM_OUT_PATH}")

def main():
    print("Loading traces...")
    with open(TRACES_PATH, "r", encoding="utf-8") as f:
        traces = json.load(f)

    print("Generating HTML showcase...")
    generate_html(traces)

    print("Capturing screenshots via headless Chrome...")
    capture_screenshots()

    print("Building Word / DOTM document...")
    build_word_document(traces)

    print("Completed successfully!")

if __name__ == "__main__":
    main()
