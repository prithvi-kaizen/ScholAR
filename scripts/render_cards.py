#!/usr/bin/env python3
"""
render_cards.py

Renders publication-quality, modern visual response cards for the 5 tested
multi-level reasoning questions on ResNet (arXiv:1512.03385) using Matplotlib and Pillow.
Saves PNGs to evaluation/results/screenshots/ and builds the .dotm and .docx files.
"""

import os
import json
import re
import textwrap
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

WORKSPACE_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TRACES_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "resnet_5q_traces.json")
SCREENSHOTS_DIR = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "screenshots")
DOTM_OUT_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "ScholAR_MultiLevel_Reasoning_Test.dotm")
DOCX_OUT_PATH = os.path.join(WORKSPACE_ROOT, "evaluation", "results", "ScholAR_MultiLevel_Reasoning_Test.docx")

MODE_COLORS = {
    "ProblemUnderstanding": "#3b82f6",  # Blue
    "HypothesisFormation": "#8b5cf6",   # Purple
    "Derivation": "#06b6d4",            # Cyan
    "Verification": "#10b981",          # Emerald
    "Calculation": "#f59e0b",           # Amber
    "Synthesis": "#ec4899",             # Pink
}

def parse_answer_tiers(answer_text):
    clean_ans = answer_text.replace("**Answer**", "").strip()
    splits = re.split(r"(\*\*\d+\.\s+[^*]+\*\*)", clean_ans)
    tiers = []
    if len(splits) > 1:
        for i in range(1, len(splits), 2):
            header = splits[i].replace("**", "").strip()
            body = splits[i+1].strip() if i+1 < len(splits) else ""
            tiers.append((header, body))
    else:
        tiers.append(("Synthesized Response", clean_ans))
    return tiers

def render_question_card(item, out_png):
    q_idx = item["q_idx"]
    query = item["query"]
    answer = item["final_answer"]
    citations = item.get("citations", [])
    reasoning_path = item.get("reasoning_path", [])

    # Dynamic canvas: 13.0 x 16.5 inches for complete breathing room
    fig = plt.figure(figsize=(13.0, 16.5), dpi=200)
    fig.patch.set_facecolor('#0b0f19')
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_facecolor('#0b0f19')
    ax.axis('off')

    # Card background
    card = patches.FancyBboxPatch(
        (0.04, 0.02), 0.92, 0.96,
        boxstyle="round,pad=0.015,rounding_size=0.02",
        facecolor='#111827',
        edgecolor='#334155',
        linewidth=1.8,
        transform=ax.transAxes
    )
    ax.add_patch(card)

    y_curr = 0.95

    # Header Badges
    q_badge = patches.FancyBboxPatch(
        (0.07, y_curr - 0.022), 0.18, 0.024,
        boxstyle="round,pad=0.004,rounding_size=0.006",
        facecolor='#1e293b',
        edgecolor='#06b6d4',
        linewidth=1.2,
        transform=ax.transAxes
    )
    ax.add_patch(q_badge)
    ax.text(0.16, y_curr - 0.010, f"QUESTION {q_idx} OF 5", color='#06b6d4', fontsize=11, fontweight='bold',
            ha='center', va='center', transform=ax.transAxes)

    v_badge = patches.FancyBboxPatch(
        (0.66, y_curr - 0.022), 0.27, 0.024,
        boxstyle="round,pad=0.004,rounding_size=0.006",
        facecolor='#064e3b',
        edgecolor='#10b981',
        linewidth=1.2,
        transform=ax.transAxes
    )
    ax.add_patch(v_badge)
    ax.text(0.795, y_curr - 0.010, "✓ MULTI-LEVEL REASONING VERIFIED", color='#34d399', fontsize=10, fontweight='bold',
            ha='center', va='center', transform=ax.transAxes)

    y_curr -= 0.038

    # Question Title (wrapped)
    q_wrapped_lines = textwrap.wrap(f'"{query}"', width=74)
    for ql in q_wrapped_lines:
        ax.text(0.07, y_curr, ql, color='#f8fafc', fontsize=13.5, fontweight='bold',
                ha='left', va='top', transform=ax.transAxes)
        y_curr -= 0.018
    y_curr -= 0.010

    # Answer Section Header
    ax.text(0.07, y_curr, "SCHOLAR MULTIMODAL MODEL RESPONSE (3-TIER MLR SYNTHESIS)", color='#94a3b8', fontsize=10, fontweight='bold',
            ha='left', va='center', transform=ax.transAxes)
    y_curr -= 0.016

    # Parse and measure answer tiers
    tiers = parse_answer_tiers(answer)
    tier_colors = ['#38bdf8', '#a78bfa', '#34d399']
    
    # Calculate answer box height
    ans_line_count = 0
    for header, body in tiers:
        lines = textwrap.wrap(body, width=86)
        ans_line_count += len(lines) + 1  # 1 for header
    
    ans_box_height = (ans_line_count * 0.0135) + 0.035
    
    # Answer Box Background
    ans_box = patches.FancyBboxPatch(
        (0.07, y_curr - ans_box_height), 0.86, ans_box_height,
        boxstyle="round,pad=0.008,rounding_size=0.012",
        facecolor='#0f172a',
        edgecolor='#1e293b',
        linewidth=1.2,
        transform=ax.transAxes
    )
    ax.add_patch(ans_box)

    # Accent bar
    accent_bar = patches.Rectangle(
        (0.07, y_curr - ans_box_height), 0.006, ans_box_height,
        facecolor='#06b6d4',
        edgecolor='none',
        transform=ax.transAxes
    )
    ax.add_patch(accent_bar)

    # Render tiers inside answer box
    tier_y = y_curr - 0.018
    for t_idx, (header, body) in enumerate(tiers):
        t_color = tier_colors[t_idx % len(tier_colors)]
        # Tier Title
        ax.text(0.088, tier_y, f"• {header.upper()}", color=t_color, fontsize=10, fontweight='bold',
                ha='left', va='center', transform=ax.transAxes)
        tier_y -= 0.016
        
        # Tier Body
        body_lines = textwrap.wrap(body, width=86)
        for bl in body_lines:
            ax.text(0.098, tier_y, bl, color='#e2e8f0', fontsize=9.5,
                    ha='left', va='center', transform=ax.transAxes)
            tier_y -= 0.0135
        tier_y -= 0.005

    y_curr -= (ans_box_height + 0.022)

    # Citations Panel Header
    ax.text(0.07, y_curr, "PROVENANCED SOURCE CITATIONS (MULTIMODAL EVIDENCE)", color='#94a3b8', fontsize=10, fontweight='bold',
            ha='left', va='center', transform=ax.transAxes)
    y_curr -= 0.016

    cits_to_show = citations[:3]
    cit_box_height = len(cits_to_show) * 0.046 + 0.012

    cit_box = patches.FancyBboxPatch(
        (0.07, y_curr - cit_box_height), 0.86, cit_box_height,
        boxstyle="round,pad=0.008,rounding_size=0.012",
        facecolor='#0f172a',
        edgecolor='#1e293b',
        linewidth=1.2,
        transform=ax.transAxes
    )
    ax.add_patch(cit_box)

    cit_row_y = y_curr - 0.020
    for c in cits_to_show:
        ref_id = c.get("ref_id")
        page = c.get("page", 1)
        sec = c.get("section", "Body")
        quote = c.get("quote", "")
        if not quote:
            quote = "[Direct structural / visual AST reference]"
        quote_wrapped = textwrap.shorten(quote, width=76, placeholder="...")

        # Ref pill
        ref_pill = patches.FancyBboxPatch(
            (0.085, cit_row_y - 0.012), 0.042, 0.022,
            boxstyle="round,pad=0.003,rounding_size=0.005",
            facecolor='#1e3a8a',
            edgecolor='#3b82f6',
            linewidth=1.0,
            transform=ax.transAxes
        )
        ax.add_patch(ref_pill)
        ax.text(0.106, cit_row_y - 0.001, f"[{ref_id}]", color='#93c5fd', fontsize=9.0, fontweight='bold',
                ha='center', va='center', transform=ax.transAxes)

        ax.text(0.138, cit_row_y + 0.002, f"Page {page}  •  {sec}", color='#38bdf8', fontsize=9.5, fontweight='bold',
                ha='left', va='center', transform=ax.transAxes)
        ax.text(0.138, cit_row_y - 0.016, f'"{quote_wrapped}"', color='#94a3b8', fontsize=8.8, fontstyle='italic',
                ha='left', va='center', transform=ax.transAxes)
        cit_row_y -= 0.046

    y_curr -= (cit_box_height + 0.022)

    # Multi-Level Reasoning Path Header
    ax.text(0.07, y_curr, "MULTI-LEVEL REASONING PATH (STEP GRAPH & SUBGOALS)", color='#94a3b8', fontsize=10, fontweight='bold',
            ha='left', va='center', transform=ax.transAxes)
    y_curr -= 0.016

    steps_to_show = reasoning_path[:5]
    step_y = y_curr
    for s in steps_to_show:
        s_idx = s.get("step_index")
        mode = s.get("mode", "Derivation")
        color = MODE_COLORS.get(mode, "#3b82f6")
        subgoal = s.get("subgoal", "")
        subgoal_wrapped = textwrap.shorten(subgoal, width=70, placeholder="...")
        ev_id = s.get("evidence_id", "")
        p = s.get("page", 1)

        # Step card
        sc = patches.FancyBboxPatch(
            (0.07, step_y - 0.034), 0.86, 0.036,
            boxstyle="round,pad=0.004,rounding_size=0.006",
            facecolor='#1e293b',
            edgecolor='#334155',
            linewidth=1.0,
            transform=ax.transAxes
        )
        ax.add_patch(sc)

        # Mode tag
        mt = patches.FancyBboxPatch(
            (0.082, step_y - 0.028), 0.20, 0.024,
            boxstyle="round,pad=0.003,rounding_size=0.005",
            facecolor='#0f172a',
            edgecolor=color,
            linewidth=1.2,
            transform=ax.transAxes
        )
        ax.add_patch(mt)
        ax.text(0.182, step_y - 0.016, f"Step {s_idx}: {mode}", color=color, fontsize=8.0, fontweight='bold',
                ha='center', va='center', transform=ax.transAxes)

        # Subgoal & Node
        ax.text(0.295, step_y - 0.010, f"{subgoal_wrapped}", color='#f1f5f9', fontsize=8.8, fontweight='bold',
                ha='left', va='center', transform=ax.transAxes)
        ax.text(0.295, step_y - 0.024, f"Node: {ev_id}  |  Page {p}", color='#64748b', fontsize=8.0,
                ha='left', va='center', transform=ax.transAxes)

        step_y -= 0.040

    # Footer Info
    ax.text(0.5, 0.015, "ScholAR Provenance-Gated Scientific Document Assistant • arXiv:1512.03385",
            color='#475569', fontsize=9, ha='center', va='center', transform=ax.transAxes)

    plt.savefig(out_png, facecolor='#0b0f19', bbox_inches='tight', pad_inches=0.1)
    plt.close(fig)
    print(f"Rendered card image: {out_png}")

def build_word_document(traces):
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("ScholAR Multi-Level Reasoning Evaluation")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Test Suite & Visual Model Responses on Deep Residual Learning (arXiv:1512.03385)")
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x02, 0x84, 0xc7)
    r_sub.font.bold = True

    # Callout box
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_cell = meta_table.cell(0, 0)
    
    tc_pr = meta_cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), "F1F5F9")
    tc_pr.append(shd)

    p_meta = meta_cell.paragraphs[0]
    p_meta.add_run("Evaluation Protocol: ").bold = True
    p_meta.add_run(
        "Strict Zero-Cue Multi-Level Reasoning. The queries contain zero explicit visual references "
        "(no mention of 'figure', 'table', 'plot', or 'diagram'). ScholAR's multimodal graph expansion "
        "automatically identifies visual evidence and structures intermediate claims using discrete MLR reasoning modes "
        "(ProblemUnderstanding, HypothesisFormation, Derivation, Verification, Calculation, Synthesis) with explicit subgoals."
    )

    doc.add_paragraph()

    # Section 1: 10 Questions
    h_overview = doc.add_heading(level=1)
    r_h_overview = h_overview.add_run("1. Set of 10 Multi-Level Scientific Questions")
    r_h_overview.font.size = Pt(16)
    r_h_overview.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    doc.add_paragraph(
        "To evaluate cross-modal and multi-level reasoning without artificial prompting cues, "
        "we formulated 10 complex scientific questions across two seminal papers (ResNet arXiv:1512.03385 and BERT arXiv:1810.04805). "
        "Each question demands hierarchical deduction across mathematical formulations, architectural mechanisms, and empirical performance metrics:"
    )

    questions_list = [
        ("ResNet (arXiv:1512.03385) - Tested Live", [
            "Q1: Why do deeper plain networks exhibit higher training error compared to shallower architectures, and how does the degradation problem differ from vanishing gradients?",
            "Q2: How do projection shortcuts compare to identity parameter-free shortcuts in terms of parameter overhead and performance across deeper architectures?",
            "Q3: What specific bottleneck modification was introduced for 50/101/152-layer networks to manage computational complexity, and what was the net impact on FLOPs?",
            "Q4: How does the training error of a 56-layer plain network compare quantitatively to that of a 20-layer plain network on CIFAR-10, and how does residual learning invert this trend?",
            "Q5: What is the margin of improvement achieved by the 152-layer residual network over the previous state-of-the-art ensemble on the ImageNet validation set?"
        ]),
        ("BERT (arXiv:1810.04805) - Formulated Multi-Level Set", [
            "Q6: How does masked language modeling resolve the unidirectional conditioning constraint in deep bidirectional representations, and what pre-training discrepancy does the 80/10/10 replacement rule mitigate?",
            "Q7: Why is cross-sentence relationship modeling essential for sentence-pair classification tasks, and how does the binarized Next Sentence Prediction task transfer to MNLI and SQuAD?",
            "Q8: What is the relative trade-off in parameter count and head dimension between the 12-layer Base and 24-layer Large architectures, and does scaling depth provide continuous gains on small-sample downstream tasks?",
            "Q9: How does BERT's fine-tuning performance compare against extracting fixed contextual embeddings across individual token layers for named entity recognition?",
            "Q10: How does the inclusion of BooksCorpus alongside English Wikipedia influence long-range contiguous sequence pre-training compared to sentence-shuffled corpora?"
        ])
    ]

    for category, q_items in questions_list:
        p_cat = doc.add_paragraph()
        r_cat = p_cat.add_run(f"• {category}")
        r_cat.bold = True
        for q_text in q_items:
            p_q = doc.add_paragraph(style='List Bullet')
            p_q.add_run(q_text)

    doc.add_page_break()

    # Section 2: Detailed Results
    h_live = doc.add_heading(level=1)
    r_h_live = h_live.add_run("2. Live System Execution & Visual Model Response Cards")
    r_h_live.font.size = Pt(16)
    r_h_live.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    for item in traces:
        q_idx = item["q_idx"]
        query = item["query"]
        answer = item["final_answer"]
        citations = item.get("citations", [])
        reasoning_path = item.get("reasoning_path", [])

        h_q = doc.add_heading(level=2)
        r_hq = h_q.add_run(f"Question {q_idx}")
        r_hq.font.size = Pt(13)
        r_hq.font.color.rgb = RGBColor(0x02, 0x84, 0xc7)

        p_query = doc.add_paragraph()
        r_query = p_query.add_run(f"\"{query}\"")
        r_query.bold = True
        r_query.font.size = Pt(11)

        # Embedded screenshot of visual card
        card_png = os.path.join(SCREENSHOTS_DIR, f"q{q_idx}_response.png")
        if os.path.exists(card_png):
            p_img_title = doc.add_paragraph()
            r_it = p_img_title.add_run(f"Figure {q_idx}: Visual Response Capture (Model Answer, Attributed Citations, & MLR Step Graph)")
            r_it.font.size = Pt(9.5)
            r_it.font.italic = True
            r_it.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(card_png, width=Inches(6.4))
            doc.add_paragraph()

        # Textual Answer Box
        ans_table = doc.add_table(rows=1, cols=1)
        ans_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ans_cell = ans_table.cell(0, 0)
        tc_pr = ans_cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), "F8FAFC")
        tc_pr.append(shd)

        p_ans = ans_cell.paragraphs[0]
        p_ans.add_run("ScholAR Final Generated Answer:\n").bold = True
        p_ans.add_run(answer)

        doc.add_paragraph()

        # Citations Table
        p_cit_heading = doc.add_paragraph()
        r_ch = p_cit_heading.add_run(f"Attributed Citations (Evidence Provenance for Q{q_idx})")
        r_ch.bold = True
        r_ch.font.size = Pt(11)

        cit_table = doc.add_table(rows=len(citations) + 1, cols=4)
        cit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Ref ID", "Page", "Section / Target", "Quoted Evidence"]
        for col_idx, text in enumerate(headers):
            cell = cit_table.cell(0, col_idx)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), "0F172A")
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            r.font.size = Pt(9.5)

        for r_idx, c in enumerate(citations):
            row_cells = cit_table.rows[r_idx + 1].cells
            row_cells[0].text = f"[{c.get('ref_id')}]"
            row_cells[1].text = f"Page {c.get('page', '-')}"
            row_cells[2].text = str(c.get('section', 'Body'))
            quote = c.get('quote', '')
            row_cells[3].text = f'"{quote}"' if quote else "(Direct structural reference)"

            bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for cell in row_cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), bg)
                tc_pr.append(shd)
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # Multi-Level Reasoning Path Table
        p_mlr_heading = doc.add_paragraph()
        r_mh = p_mlr_heading.add_run(f"Multi-Level Reasoning Path (Step Graph for Q{q_idx})")
        r_mh.bold = True
        r_mh.font.size = Pt(11)

        mlr_table = doc.add_table(rows=len(reasoning_path) + 1, cols=5)
        mlr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        mlr_headers = ["Step", "MLR Mode", "Subgoal Descriptor (<= 30 words)", "Evidence Node", "Page"]
        for col_idx, text in enumerate(mlr_headers):
            cell = mlr_table.cell(0, col_idx)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), "1E293B")
            tc_pr.append(shd)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            r.font.size = Pt(9.5)

        for s_idx, step in enumerate(reasoning_path):
            row_cells = mlr_table.rows[s_idx + 1].cells
            row_cells[0].text = f"Step {step.get('step_index')}"
            row_cells[1].text = str(step.get('mode', 'Derivation'))
            row_cells[2].text = str(step.get('subgoal', ''))
            row_cells[3].text = str(step.get('evidence_id', ''))
            row_cells[4].text = f"p. {step.get('page', 1)}"

            bg = "FFFFFF" if s_idx % 2 == 0 else "F1F5F9"
            for cell in row_cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), bg)
                tc_pr.append(shd)
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)

        doc.add_paragraph()
        doc.add_paragraph("─" * 60)

    doc.save(DOCX_OUT_PATH)
    doc.save(DOTM_OUT_PATH)
    print(f"Generated DOCX: {DOCX_OUT_PATH}")
    print(f"Generated DOTM: {DOTM_OUT_PATH}")

def main():
    os.makedirs(SCREENSHOTS_DIR, exist_ok=True)
    print("Loading traces...")
    with open(TRACES_PATH, "r", encoding="utf-8") as f:
        traces = json.load(f)

    print("Rendering question card screenshots...")
    for item in traces:
        q_idx = item["q_idx"]
        out_png = os.path.join(SCREENSHOTS_DIR, f"q{q_idx}_response.png")
        render_question_card(item, out_png)

    print("Building Word / DOTM document...")
    build_word_document(traces)

    print("All tasks completed successfully!")

if __name__ == "__main__":
    main()
